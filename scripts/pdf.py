import re
from docx import Document


def get_text(filename):
    pass

def docx_replace_regex(doc_obj, regex, replace):
    for p in doc_obj.paragraphs:
        if regex.search(p.text):
            inline = p.runs
            # Loop added to work with runs (strings with same style)
            for i in range(len(inline)):
                if regex.search(inline[i].text):
                    text = regex.sub(replace, inline[i].text)
                    inline[i].text = text

    for table in doc_obj.tables:
        for row in table.rows:
            for cell in row.cells:
                docx_replace_regex(cell, regex, replace)


regex1 = re.compile(r"your regex")
replace1 = r"your replace string"
filename = ".docx"
doc = Document(filename)
docx_replace_regex(doc, regex1, replace1)
doc.save("result1.docx")
